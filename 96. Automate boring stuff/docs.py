#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sat Dec 19 14:21:30 2020

@author: yukuo
"""

import docx
import os
os.chdir('/Users/yukuo/Desktop/practice/96. Automate boring stuff')

d = docx.Document()  # create document 
d.add_paragraph('你好，oh hi yo u.')
d.add_paragraph('hello!! /n')

p = d.paragraphs[0]
p.add_run('new run!!')
p.runs[1].bold = True
p.runs[1].italic = True

d.save('doc2.docx')
